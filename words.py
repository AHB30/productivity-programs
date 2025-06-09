import random
from docx import Document

def load_words_from_docx(filename="Words.docx"):
    """
    Load words from a Word file.
    Each paragraph is treated as a separate word.
    """
    try:
        doc = Document(filename)
        words = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    except Exception as e:
        print(f"[ERROR] Failed to load document: {e}")
        words = []
    return words


def save_words_to_docx(words, filename="Words.docx"):
    """
    Save the remaining words back into the Word document.
    """
    try:
        doc = Document()
        for word in words:
            doc.add_paragraph(word)
        doc.save(filename)
    except Exception as e:
        print(f"[ERROR] Failed to save document: {e}")


def pick_random_word(words):
    """
    Select a random word from the list and remove it.
    """
    if not words:
        return None
    word = random.choice(words)
    words.remove(word)
    return word


def main():
    words = load_words_from_docx()
    if not words:
        print("[INFO] No words available in the document.")
        return

    word = pick_random_word(words)
    save_words_to_docx(words)

    if word:
        print(f"Selected word: {word}")
    else:
        print("[INFO] No words left in the list!")


if __name__ == "__main__":
    main()
